"""
Report generation system for monitoring results
"""

import logging
import json
from datetime import datetime, timedelta
from pathlib import Path
from typing import List, Dict, Any, Optional
from jinja2 import Template, Environment, FileSystemLoader
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

from core.monitor import MonitorResult
from core.variables import VariableManager
from config.settings import config_manager
from utils.helpers import ensure_directory, format_size


class ReportGenerator:
    """Generates monitoring reports in various formats"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        self.config = config_manager.get_reports_config()
        self.directories_config = config_manager.get_directories_config()
        
        # Setup reports directory
        self.reports_dir = ensure_directory(self.directories_config.reports)
        self.templates_dir = Path(__file__).parent / "templates"
        self.templates_dir.mkdir(exist_ok=True)
        
        # Setup Jinja2 environment
        self.jinja_env = Environment(
            loader=FileSystemLoader(str(self.templates_dir)),
            autoescape=True
        )
        
        # Create default templates if they don't exist
        self.create_default_templates()
    
    def create_default_templates(self):
        """Create default report templates"""
        # HTML template
        html_template_path = self.templates_dir / "default.html"
        if not html_template_path.exists():
            self.create_default_html_template(html_template_path)
        
        # Email template
        email_template_path = self.templates_dir / "email.html"
        if not email_template_path.exists():
            self.create_default_email_template(email_template_path)
    
    def create_default_html_template(self, template_path: Path):
        """Create default HTML report template"""
        template_content = '''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Website Monitoring Report</title>
    <style>
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            margin: 0;
            padding: 20px;
            background-color: #f5f5f5;
        }
        .container {
            max-width: 1200px;
            margin: 0 auto;
            background: white;
            border-radius: 8px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
            overflow: hidden;
        }
        .header {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 30px;
            text-align: center;
        }
        .header h1 {
            margin: 0;
            font-size: 2.5em;
            font-weight: 300;
        }
        .header .subtitle {
            margin-top: 10px;
            opacity: 0.9;
            font-size: 1.1em;
        }
        .summary {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(250px, 1fr));
            gap: 20px;
            padding: 30px;
            background-color: #f8f9fa;
        }
        .summary-card {
            background: white;
            padding: 20px;
            border-radius: 8px;
            text-align: center;
            box-shadow: 0 2px 5px rgba(0,0,0,0.1);
        }
        .summary-card h3 {
            margin: 0 0 10px 0;
            color: #333;
            font-size: 1.1em;
        }
        .summary-card .value {
            font-size: 2em;
            font-weight: bold;
            margin: 10px 0;
        }
        .success { color: #28a745; }
        .warning { color: #ffc107; }
        .danger { color: #dc3545; }
        .info { color: #17a2b8; }
        
        .results {
            padding: 30px;
        }
        .results h2 {
            margin-bottom: 20px;
            color: #333;
            border-bottom: 2px solid #667eea;
            padding-bottom: 10px;
        }
        .result-item {
            background: white;
            border: 1px solid #ddd;
            border-radius: 8px;
            margin-bottom: 20px;
            overflow: hidden;
        }
        .result-header {
            padding: 15px 20px;
            background-color: #f8f9fa;
            border-bottom: 1px solid #ddd;
            display: flex;
            justify-content: space-between;
            align-items: center;
        }
        .result-header.success {
            background-color: #d4edda;
            border-color: #c3e6cb;
        }
        .result-header.failed {
            background-color: #f8d7da;
            border-color: #f5c6cb;
        }
        .result-body {
            padding: 20px;
        }
        .result-details {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 15px;
            margin-bottom: 15px;
        }
        .detail-item {
            display: flex;
            justify-content: space-between;
            padding: 8px 0;
            border-bottom: 1px solid #eee;
        }
        .detail-label {
            font-weight: bold;
            color: #666;
        }
        .screenshot {
            text-align: center;
            margin-top: 20px;
        }
        .screenshot img {
            max-width: 100%;
            height: auto;
            border: 1px solid #ddd;
            border-radius: 4px;
        }
        .variables {
            margin-top: 20px;
            padding: 15px;
            background-color: #f8f9fa;
            border-radius: 4px;
        }
        .variables h4 {
            margin: 0 0 10px 0;
            color: #333;
        }
        .variable-item {
            margin: 5px 0;
            padding: 5px 10px;
            background: white;
            border-radius: 3px;
            font-family: monospace;
            font-size: 0.9em;
        }
        .footer {
            background-color: #333;
            color: white;
            text-align: center;
            padding: 20px;
            font-size: 0.9em;
        }
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>Website Monitoring Report</h1>
            <div class="subtitle">Generated on {{ report_date }}</div>
        </div>
        
        <div class="summary">
            <div class="summary-card">
                <h3>Total Checks</h3>
                <div class="value info">{{ summary.total_checks }}</div>
            </div>
            <div class="summary-card">
                <h3>Successful</h3>
                <div class="value success">{{ summary.successful_checks }}</div>
            </div>
            <div class="summary-card">
                <h3>Failed</h3>
                <div class="value danger">{{ summary.failed_checks }}</div>
            </div>
            <div class="summary-card">
                <h3>Success Rate</h3>
                <div class="value {% if summary.success_rate > 95 %}success{% elif summary.success_rate > 80 %}warning{% else %}danger{% endif %}">
                    {{ "%.1f"|format(summary.success_rate) }}%
                </div>
            </div>
        </div>
        
        <div class="results">
            <h2>Detailed Results</h2>
            {% for result in results %}
            <div class="result-item">
                <div class="result-header {% if result.success %}success{% else %}failed{% endif %}">
                    <div>
                        <strong>{{ result.website_name }}</strong>
                        <span style="margin-left: 10px; font-size: 0.9em; color: #666;">
                            {{ result.timestamp.strftime('%Y-%m-%d %H:%M:%S') }}
                        </span>
                    </div>
                    <div>
                        {% if result.success %}
                            <span class="success">✓ SUCCESS</span>
                        {% else %}
                            <span class="danger">✗ FAILED</span>
                        {% endif %}
                    </div>
                </div>
                <div class="result-body">
                    <div class="result-details">
                        <div class="detail-item">
                            <span class="detail-label">URL:</span>
                            <span>{{ result.url }}</span>
                        </div>
                        <div class="detail-item">
                            <span class="detail-label">Status Code:</span>
                            <span>{{ result.status_code or 'N/A' }}</span>
                        </div>
                        <div class="detail-item">
                            <span class="detail-label">Response Time:</span>
                            <span>{{ "%.2f"|format(result.response_time or 0) }}s</span>
                        </div>
                        <div class="detail-item">
                            <span class="detail-label">Content Size:</span>
                            <span>{{ format_size(result.content_size or 0) }}</span>
                        </div>
                    </div>
                    
                    {% if result.error_message %}
                    <div style="color: #dc3545; margin: 10px 0; padding: 10px; background-color: #f8d7da; border-radius: 4px;">
                        <strong>Error:</strong> {{ result.error_message }}
                    </div>
                    {% endif %}
                    
                    {% if result.screenshot_path %}
                    <div class="screenshot">
                        <h4>Screenshot</h4>
                        <img src="{{ result.screenshot_path }}" alt="Website Screenshot" />
                    </div>
                    {% endif %}
                    
                    {% if result.variables %}
                    <div class="variables">
                        <h4>Variables</h4>
                        {% for var_name, var_value in result.variables.items() %}
                        <div class="variable-item">
                            <strong>{{ var_name }}:</strong> {{ var_value|string|truncate(100) }}
                        </div>
                        {% endfor %}
                    </div>
                    {% endif %}
                </div>
            </div>
            {% endfor %}
        </div>
        
        <div class="footer">
            <p>Generated by Website Monitoring and Reporting System</p>
            <p>Total execution time: {{ summary.total_execution_time }}s</p>
        </div>
    </div>
</body>
</html>'''
        
        with open(template_path, 'w', encoding='utf-8') as f:
            f.write(template_content)
        
        self.logger.info(f"Created default HTML template: {template_path}")
    
    def create_default_email_template(self, template_path: Path):
        """Create default email template"""
        template_content = '''<div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
    <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 20px; text-align: center;">
        <h1 style="margin: 0;">Monitoring Alert</h1>
        <p style="margin: 10px 0 0 0;">{{ alert_type }} - {{ timestamp }}</p>
    </div>
    
    <div style="padding: 20px; background-color: #f9f9f9;">
        <h2 style="color: #333; margin-top: 0;">Summary</h2>
        <p><strong>Website:</strong> {{ website_name }}</p>
        <p><strong>URL:</strong> {{ url }}</p>
        <p><strong>Status:</strong> 
            <span style="color: {% if success %}#28a745{% else %}#dc3545{% endif %};">
                {% if success %}✓ SUCCESS{% else %}✗ FAILED{% endif %}
            </span>
        </p>
        
        {% if error_message %}
        <div style="background-color: #f8d7da; border: 1px solid #f5c6cb; padding: 10px; border-radius: 4px; margin: 10px 0;">
            <strong>Error:</strong> {{ error_message }}
        </div>
        {% endif %}
        
        <h3 style="color: #333;">Details</h3>
        <ul style="list-style: none; padding: 0;">
            <li style="margin: 5px 0;"><strong>Response Time:</strong> {{ "%.2f"|format(response_time or 0) }}s</li>
            <li style="margin: 5px 0;"><strong>Status Code:</strong> {{ status_code or 'N/A' }}</li>
            <li style="margin: 5px 0;"><strong>Timestamp:</strong> {{ timestamp }}</li>
        </ul>
    </div>
    
    <div style="background-color: #333; color: white; text-align: center; padding: 15px; font-size: 0.9em;">
        <p style="margin: 0;">Website Monitoring and Reporting System</p>
    </div>
</div>'''
        
        with open(template_path, 'w', encoding='utf-8') as f:
            f.write(template_content)
        
        self.logger.info(f"Created default email template: {template_path}")
    
    def generate_html_report(self, results: List[MonitorResult], template_name: str = "default", 
                           variables: Optional[Dict[str, Any]] = None) -> str:
        """Generate HTML report from monitoring results"""
        try:
            # Calculate summary statistics
            summary = self._calculate_summary(results)
            
            # Prepare template variables
            template_vars = {
                'report_date': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                'results': results,
                'summary': summary,
                'variables': variables or {},
                'format_size': format_size
            }
            
            # Load and render template
            template = self.jinja_env.get_template(f"{template_name}.html")
            html_content = template.render(**template_vars)
            
            # Generate filename and save report
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"monitoring_report_{timestamp}.html"
            report_path = self.reports_dir / filename
            
            with open(report_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            self.logger.info(f"Generated HTML report: {report_path}")
            return str(report_path)
            
        except Exception as e:
            self.logger.error(f"Failed to generate HTML report: {e}")
            raise
    
    def generate_json_report(self, results: List[MonitorResult], include_summary: bool = True) -> str:
        """Generate JSON report from monitoring results"""
        try:
            # Prepare data
            report_data = {
                'generated_at': datetime.now().isoformat(),
                'results': []
            }
            
            # Add summary if requested
            if include_summary:
                report_data['summary'] = self._calculate_summary(results)
            
            # Convert results to dict format
            for result in results:
                result_dict = {
                    'website_name': result.website_name,
                    'url': result.url,
                    'timestamp': result.timestamp.isoformat(),
                    'success': result.success,
                    'status_code': result.status_code,
                    'response_time': result.response_time,
                    'error_message': result.error_message,
                    'content_size': result.content_size,
                    'screenshot_path': result.screenshot_path,
                    'custom_check_results': result.custom_check_results,
                    'downloaded_files': result.downloaded_files,
                    'variables': result.variables
                }
                report_data['results'].append(result_dict)
            
            # Generate filename and save report
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"monitoring_report_{timestamp}.json"
            report_path = self.reports_dir / filename
            
            with open(report_path, 'w', encoding='utf-8') as f:
                json.dump(report_data, f, indent=2, ensure_ascii=False)
            
            self.logger.info(f"Generated JSON report: {report_path}")
            return str(report_path)
            
        except Exception as e:
            self.logger.error(f"Failed to generate JSON report: {e}")
            raise
    
    def generate_email_content(self, result: MonitorResult, alert_type: str = "Monitoring Alert") -> str:
        """Generate email content for a monitoring result"""
        try:
            template_vars = {
                'alert_type': alert_type,
                'website_name': result.website_name,
                'url': result.url,
                'success': result.success,
                'status_code': result.status_code,
                'response_time': result.response_time,
                'error_message': result.error_message,
                'timestamp': result.timestamp.strftime('%Y-%m-%d %H:%M:%S')
            }
            
            template = self.jinja_env.get_template("email.html")
            return template.render(**template_vars)
            
        except Exception as e:
            self.logger.error(f"Failed to generate email content: {e}")
            raise
    
    def generate_patrol_word_report(self, patrol_results: List, task_name: str = "巡检任务", 
                                  template_name: str = "default") -> str:
        """Generate Word document report from patrol results with screenshots and extracted values"""
        try:
            # Import PatrolResult if needed
            from core.patrol import PatrolResult
            
            # Create a new document
            doc = Document()
            
            # Set document title
            title = doc.add_heading('网站巡检报告', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add report generation time
            report_time = doc.add_paragraph(f'报告生成时间: {datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}')
            report_time.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add task name
            task_para = doc.add_paragraph(f'巡检任务: {task_name}')
            task_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Calculate summary from patrol results
            summary = self._calculate_patrol_summary(patrol_results)
            
            # Add summary section
            doc.add_heading('巡检概要', level=1)
            summary_table = doc.add_table(rows=7, cols=2)
            summary_table.style = 'Table Grid'
            summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Populate summary table
            summary_data = [
                ('总计巡检网站', str(summary['total_websites'])),
                ('成功网站数', str(summary['successful_websites'])),
                ('失败网站数', str(summary['failed_websites'])),
                ('成功率', f"{summary['success_rate']:.1f}%"),
                ('平均响应时间', f"{summary['average_response_time']:.2f}秒"),
                ('总执行检查数', str(summary['total_checks'])),
                ('提取数据项', str(summary['extracted_values']))
            ]
            
            for i, (label, value) in enumerate(summary_data):
                row = summary_table.rows[i]
                row.cells[0].text = label
                row.cells[1].text = value
                # Make header cells bold
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if cell == row.cells[0]:  # Header column
                                run.bold = True
            
            # Add detailed results section
            doc.add_heading('详细巡检结果', level=1)
            
            for result in patrol_results:
                # Add website section
                doc.add_heading(f'{result.task_name} - {result.website_url}', level=2)
                
                # Basic info table
                info_table = doc.add_table(rows=4, cols=2)
                info_table.style = 'Table Grid'
                
                info_data = [
                    ('网址', result.website_url),
                    ('检查时间', result.timestamp.strftime("%Y年%m月%d日 %H:%M:%S")),
                    ('状态', "成功" if result.success else "失败"),
                    ('响应时间', f"{result.response_time:.2f}秒" if result.response_time else "未知")
                ]
                
                if result.status_code:
                    info_data.append(('状态码', str(result.status_code)))
                
                for i, (label, value) in enumerate(info_data):
                    if i < len(info_table.rows):
                        row = info_table.rows[i]
                        row.cells[0].text = label
                        row.cells[1].text = str(value)
                        # Make label bold
                        for paragraph in row.cells[0].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                
                # Add screenshot if available
                if result.screenshot_path and Path(result.screenshot_path).exists():
                    doc.add_paragraph("页面截图:")
                    try:
                        screenshot_paragraph = doc.add_paragraph()
                        run = screenshot_paragraph.runs[0] if screenshot_paragraph.runs else screenshot_paragraph.add_run()
                        run.add_picture(result.screenshot_path, width=Inches(5.0))
                    except Exception as e:
                        doc.add_paragraph(f"截图加载失败: {e}")
                
                # Add check results if available
                if result.check_results:
                    doc.add_heading('检查结果', level=3)
                    
                    check_table = doc.add_table(rows=1, cols=4)
                    check_table.style = 'Table Grid'
                    
                    # Header row
                    header_cells = check_table.rows[0].cells
                    headers = ['检查项', '结果', '提取值', '说明']
                    for i, header in enumerate(headers):
                        header_cells[i].text = header
                        for paragraph in header_cells[i].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                    
                    # Add check data
                    for check_name, check_result in result.check_results.items():
                        row_cells = check_table.add_row().cells
                        row_cells[0].text = check_name
                        row_cells[1].text = "✓" if check_result.get('success') else "✗"
                        row_cells[2].text = str(check_result.get('extracted_value', '')) if check_result.get('extracted_value') else '无'
                        row_cells[3].text = check_result.get('message', '')
                
                # Add error details if failed
                if not result.success and result.error_message:
                    doc.add_heading('错误详情', level=3)
                    doc.add_paragraph(result.error_message)
                
                # Add extracted data if available
                if result.extracted_data:
                    doc.add_heading('提取数据', level=3)
                    for key, value in result.extracted_data.items():
                        data_para = doc.add_paragraph()
                        data_para.add_run(f"{key}: ").bold = True
                        data_para.add_run(str(value))
                
                # Add separator
                doc.add_paragraph()
            
            # Add failed results summary if any
            failed_results = [r for r in patrol_results if not r.success]
            if failed_results:
                doc.add_heading('失败汇总', level=1)
                for result in failed_results:
                    failure_para = doc.add_paragraph()
                    failure_para.add_run(f"• {result.website_url}: ").bold = True
                    failure_para.add_run(result.error_message or '未知错误')
            
            # Generate filename and save document
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_task_name = "".join(c for c in task_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
            filename = f"patrol_report_{safe_task_name}_{timestamp}.docx"
            report_path = self.reports_dir / filename
            
            doc.save(str(report_path))
            
            self.logger.info(f"Generated patrol Word report: {report_path}")
            return str(report_path)
            
        except Exception as e:
            self.logger.error(f"Failed to generate patrol Word report: {e}")
            raise
    
    def _calculate_patrol_summary(self, patrol_results: List) -> Dict[str, Any]:
        """Calculate summary statistics from patrol results"""
        if not patrol_results:
            return {
                'total_websites': 0,
                'successful_websites': 0,
                'failed_websites': 0,
                'success_rate': 0.0,
                'average_response_time': 0.0,
                'total_checks': 0,
                'extracted_values': 0
            }
        
        total_websites = len(patrol_results)
        successful_websites = sum(1 for r in patrol_results if r.success)
        failed_websites = total_websites - successful_websites
        success_rate = (successful_websites / total_websites) * 100
        
        # Calculate average response time
        response_times = [r.response_time for r in patrol_results if r.response_time is not None]
        average_response_time = sum(response_times) / len(response_times) if response_times else 0
        
        # Count total checks performed
        total_checks = sum(len(r.check_results) for r in patrol_results if r.check_results)
        
        # Count extracted values
        extracted_values = 0
        for result in patrol_results:
            if result.check_results:
                for check_result in result.check_results.values():
                    if check_result.get('extracted_value'):
                        extracted_values += 1
            if result.extracted_data:
                extracted_values += len(result.extracted_data)
        
        return {
            'total_websites': total_websites,
            'successful_websites': successful_websites,
            'failed_websites': failed_websites,
            'success_rate': success_rate,
            'average_response_time': average_response_time,
            'total_checks': total_checks,
            'extracted_values': extracted_values
        }
    
    def _calculate_summary(self, results: List[MonitorResult]) -> Dict[str, Any]:
        """Calculate summary statistics from results"""
        if not results:
            return {
                'total_checks': 0,
                'successful_checks': 0,
                'failed_checks': 0,
                'success_rate': 0.0,
                'average_response_time': 0.0,
                'total_execution_time': 0.0
            }
        
        total_checks = len(results)
        successful_checks = sum(1 for r in results if r.success)
        failed_checks = total_checks - successful_checks
        success_rate = (successful_checks / total_checks) * 100 if total_checks > 0 else 0
        
        # Calculate average response time (excluding None values)
        response_times = [r.response_time for r in results if r.response_time is not None]
        avg_response_time = sum(response_times) / len(response_times) if response_times else 0
        
        # Calculate total execution time
        total_execution_time = sum(response_times) if response_times else 0
        
        return {
            'total_checks': total_checks,
            'successful_checks': successful_checks,
            'failed_checks': failed_checks,
            'success_rate': success_rate,
            'average_response_time': avg_response_time,
            'total_execution_time': total_execution_time
        }
    
    def cleanup_old_reports(self, days: int = 30) -> None:
        """Remove old report files"""
        import time
        
        current_time = time.time()
        cutoff_time = current_time - (days * 24 * 3600)
        
        for report_file in self.reports_dir.glob("*"):
            if report_file.is_file():
                file_time = report_file.stat().st_mtime
                if file_time < cutoff_time:
                    try:
                        report_file.unlink()
                        self.logger.info(f"Deleted old report: {report_file}")
                    except Exception as e:
                        self.logger.error(f"Failed to delete report {report_file}: {e}")
    
    def get_available_templates(self) -> List[str]:
        """Get list of available report templates"""
        templates = []
        for template_file in self.templates_dir.glob("*.html"):
            if template_file.name != "email.html":  # Exclude email template
                templates.append(template_file.stem)
        return templates