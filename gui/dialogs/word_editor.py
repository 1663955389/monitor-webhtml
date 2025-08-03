"""
Word report editing dialog for modifying generated reports
"""

import logging
import os
from datetime import datetime
from pathlib import Path
from typing import Optional

from PyQt5.QtWidgets import (
    QDialog, QVBoxLayout, QHBoxLayout, QPushButton, QTextEdit,
    QLabel, QFileDialog, QMessageBox, QSplitter, QGroupBox,
    QTableWidget, QTableWidgetItem, QHeaderView, QComboBox,
    QListWidget, QListWidgetItem
)
from PyQt5.QtCore import Qt, pyqtSignal
from PyQt5.QtGui import QFont, QBrush, QColor

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from core.variables import VariableManager


class WordReportEditor(QDialog):
    """Dialog for editing Word reports with preview and modification capabilities"""
    
    report_saved = pyqtSignal(str)  # Emit path when report is saved
    
    def __init__(self, parent=None, report_path: str = None):
        super().__init__(parent)
        self.logger = logging.getLogger(__name__)
        self.current_report_path = report_path
        self.document = None
        self.variable_manager = VariableManager()
        
        self.init_ui()
        
        if report_path and Path(report_path).exists():
            self.load_report(report_path)
    
    def init_ui(self):
        """Initialize the user interface"""
        self.setWindowTitle("Word报告编辑器")
        self.setGeometry(100, 100, 1000, 700)
        
        # Main layout
        main_layout = QVBoxLayout(self)
        
        # Top toolbar
        toolbar_layout = QHBoxLayout()
        
        self.open_btn = QPushButton("打开报告")
        self.open_btn.clicked.connect(self.open_report)
        toolbar_layout.addWidget(self.open_btn)
        
        self.new_btn = QPushButton("新建报告")
        self.new_btn.clicked.connect(self.new_report)
        toolbar_layout.addWidget(self.new_btn)
        
        # Template dropdown for new reports
        self.template_combo = QComboBox()
        self.template_combo.addItems(["基础模板", "详细巡检模板", "简要汇总模板"])
        self.template_combo.setToolTip("选择新建报告的模板类型")
        toolbar_layout.addWidget(QLabel("模板:"))
        toolbar_layout.addWidget(self.template_combo)
        
        self.save_btn = QPushButton("保存报告")
        self.save_btn.clicked.connect(self.save_report)
        self.save_btn.setEnabled(False)
        toolbar_layout.addWidget(self.save_btn)
        
        self.save_as_btn = QPushButton("另存为")
        self.save_as_btn.clicked.connect(self.save_as_report)
        self.save_as_btn.setEnabled(False)
        toolbar_layout.addWidget(self.save_as_btn)
        
        toolbar_layout.addStretch()
        
        self.preview_btn = QPushButton("刷新预览")
        self.preview_btn.clicked.connect(self.refresh_preview)
        self.preview_btn.setEnabled(False)
        toolbar_layout.addWidget(self.preview_btn)
        
        main_layout.addLayout(toolbar_layout)
        
        # Main content area with splitter
        splitter = QSplitter(Qt.Horizontal)
        main_layout.addWidget(splitter)
        
        # Left side: Structure and editing
        left_widget = QGroupBox("报告结构编辑")
        left_layout = QVBoxLayout(left_widget)
        
        # Document structure
        structure_label = QLabel("文档结构:")
        left_layout.addWidget(structure_label)
        
        self.structure_table = QTableWidget()
        self.structure_table.setColumnCount(4)
        self.structure_table.setHorizontalHeaderLabels(["类型", "内容", "编辑", "删除"])
        self.structure_table.horizontalHeader().setStretchLastSection(False)
        self.structure_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeToContents)
        self.structure_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.Stretch)
        self.structure_table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeToContents)
        self.structure_table.horizontalHeader().setSectionResizeMode(3, QHeaderView.ResizeToContents)
        left_layout.addWidget(self.structure_table)
        
        # Content editor
        editor_label = QLabel("内容编辑 (支持变量: ${变量名}):")
        left_layout.addWidget(editor_label)
        
        self.content_editor = QTextEdit()
        self.content_editor.setFont(QFont("Arial", 10))
        self.content_editor.setPlaceholderText("选择要编辑的段落或表格...\n支持变量语法: ${变量名}")
        left_layout.addWidget(self.content_editor)
        
        # Edit controls
        edit_controls_layout = QHBoxLayout()
        
        self.update_content_btn = QPushButton("更新内容")
        self.update_content_btn.clicked.connect(self.update_selected_content)
        self.update_content_btn.setEnabled(False)
        edit_controls_layout.addWidget(self.update_content_btn)
        
        self.add_paragraph_btn = QPushButton("添加段落")
        self.add_paragraph_btn.clicked.connect(self.add_paragraph)
        self.add_paragraph_btn.setEnabled(False)
        edit_controls_layout.addWidget(self.add_paragraph_btn)
        
        self.insert_variable_btn = QPushButton("插入变量")
        self.insert_variable_btn.clicked.connect(self.show_variable_selector)
        self.insert_variable_btn.setEnabled(True)
        edit_controls_layout.addWidget(self.insert_variable_btn)
        
        edit_controls_layout.addStretch()
        left_layout.addLayout(edit_controls_layout)
        
        splitter.addWidget(left_widget)
        
        # Right side: Preview and Variables
        right_widget = QGroupBox("预览和变量")
        right_layout = QVBoxLayout(right_widget)
        
        # Available variables section
        variables_label = QLabel("可用变量:")
        right_layout.addWidget(variables_label)
        
        self.variables_list = QListWidget()
        self.variables_list.setMaximumHeight(150)
        self.variables_list.itemDoubleClicked.connect(self.insert_selected_variable)
        right_layout.addWidget(self.variables_list)
        
        # Preview section
        preview_label = QLabel("文档预览:")
        right_layout.addWidget(preview_label)
        
        self.preview_text = QTextEdit()
        self.preview_text.setReadOnly(True)
        self.preview_text.setFont(QFont("Consolas", 9))
        self.preview_text.setPlaceholderText("文档预览将在这里显示...")
        right_layout.addWidget(self.preview_text)
        
        splitter.addWidget(right_widget)
        
        # Set splitter proportions
        splitter.setStretchFactor(0, 1)
        splitter.setStretchFactor(1, 1)
        
        # Bottom buttons
        button_layout = QHBoxLayout()
        button_layout.addStretch()
        
        self.close_btn = QPushButton("关闭")
        self.close_btn.clicked.connect(self.accept)
        button_layout.addWidget(self.close_btn)
        
        main_layout.addLayout(button_layout)
        
        # Connect table selection
        self.structure_table.currentCellChanged.connect(self.on_structure_selection_changed)
        
        # Load available variables
        self.refresh_variables_list()
    
    def open_report(self):
        """Open an existing Word report"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, "选择Word报告", "", "Word文档 (*.docx);;所有文件 (*.*)"
        )
        
        if file_path:
            self.load_report(file_path)
    
    def new_report(self):
        """Create a new blank report"""
        try:
            # Create a new document
            self.document = Document()
            self.current_report_path = None
            
            # Get selected template
            template_type = self.template_combo.currentText()
            
            # Add content based on template
            if template_type == "基础模板":
                self._create_basic_template()
            elif template_type == "详细巡检模板":
                self._create_detailed_patrol_template()
            elif template_type == "简要汇总模板":
                self._create_summary_template()
            
            # Update UI state
            self.save_btn.setEnabled(False)  # Disable save until we have a path
            self.save_as_btn.setEnabled(True)
            self.preview_btn.setEnabled(True)
            self.add_paragraph_btn.setEnabled(True)
            
            # Populate structure table
            self.populate_structure_table()
            
            # Refresh preview
            self.refresh_preview()
            
            self.logger.info(f"Created new report with template: {template_type}")
            
        except Exception as e:
            QMessageBox.critical(self, "错误", f"无法创建新报告:\n{str(e)}")
            self.logger.error(f"Failed to create new report: {e}")
    
    def _create_basic_template(self):
        """Create basic report template"""
        title = self.document.add_heading('自定义报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.document.add_paragraph(f'创建时间: {datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}')
        
        self.document.add_heading('概述', level=1)
        self.document.add_paragraph('在此编写报告概述...')
        
        self.document.add_heading('内容', level=1)
        self.document.add_paragraph('在此编写主要内容...')
    
    def _create_detailed_patrol_template(self):
        """Create detailed patrol report template"""
        title = self.document.add_heading('网站巡检详细报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Report info
        info_para = self.document.add_paragraph(f'报告生成时间: {datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}')
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Executive summary
        self.document.add_heading('执行摘要', level=1)
        self.document.add_paragraph('本次巡检的总体情况概述...')
        
        # Statistics section
        self.document.add_heading('统计概览', level=1)
        stats_table = self.document.add_table(rows=5, cols=2)
        stats_table.style = 'Table Grid'
        
        stats_data = [
            ('巡检网站总数', '[待填写]'),
            ('成功网站数', '[待填写]'),
            ('失败网站数', '[待填写]'),
            ('成功率', '[待填写]'),
            ('平均响应时间', '[待填写]')
        ]
        
        for i, (label, value) in enumerate(stats_data):
            row = stats_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
        
        # Detailed results
        self.document.add_heading('详细巡检结果', level=1)
        self.document.add_paragraph('各网站的详细巡检结果...')
        
        # Issues and recommendations
        self.document.add_heading('问题与建议', level=1)
        self.document.add_paragraph('发现的问题及改进建议...')
        
        # Conclusion
        self.document.add_heading('结论', level=1)
        self.document.add_paragraph('巡检结论与后续行动计划...')
    
    def _create_summary_template(self):
        """Create summary report template"""
        title = self.document.add_heading('巡检汇总报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.document.add_paragraph(f'报告时间: {datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}')
        
        # Quick stats
        self.document.add_heading('快速统计', level=1)
        summary_table = self.document.add_table(rows=3, cols=2)
        summary_table.style = 'Table Grid'
        
        summary_data = [
            ('总巡检数', '[待填写]'),
            ('成功率', '[待填写]'),
            ('主要问题', '[待填写]')
        ]
        
        for i, (label, value) in enumerate(summary_data):
            row = summary_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
        
        # Key findings
        self.document.add_heading('主要发现', level=1)
        self.document.add_paragraph('• 重要发现1')
        self.document.add_paragraph('• 重要发现2')
        self.document.add_paragraph('• 重要发现3')
        
        # Action items
        self.document.add_heading('行动项', level=1)
        self.document.add_paragraph('• 待办事项1')
        self.document.add_paragraph('• 待办事项2')
    
    def load_report(self, file_path: str):
        """Load a Word report for editing"""
        try:
            self.document = Document(file_path)
            self.current_report_path = file_path
            
            # Update UI state
            self.save_btn.setEnabled(True)
            self.save_as_btn.setEnabled(True)
            self.preview_btn.setEnabled(True)
            self.add_paragraph_btn.setEnabled(True)
            
            # Populate structure table
            self.populate_structure_table()
            
            # Refresh preview
            self.refresh_preview()
            
            self.logger.info(f"Loaded Word report: {file_path}")
            
        except Exception as e:
            QMessageBox.critical(self, "错误", f"无法加载报告文件:\n{str(e)}")
            self.logger.error(f"Failed to load Word report: {e}")
    
    def populate_structure_table(self):
        """Populate the structure table with document elements"""
        if not self.document:
            return
        
        self.structure_table.setRowCount(0)
        
        for i, paragraph in enumerate(self.document.paragraphs):
            row = self.structure_table.rowCount()
            self.structure_table.insertRow(row)
            
            # Determine type
            if paragraph.style.name.startswith('Heading'):
                element_type = f"标题 {paragraph.style.name[-1]}"
            else:
                element_type = "段落"
            
            # Content preview (first 50 chars)
            content = paragraph.text[:50] + "..." if len(paragraph.text) > 50 else paragraph.text
            
            self.structure_table.setItem(row, 0, QTableWidgetItem(element_type))
            self.structure_table.setItem(row, 1, QTableWidgetItem(content))
            
            # Edit button
            edit_btn = QPushButton("编辑")
            edit_btn.clicked.connect(lambda checked, idx=i: self.edit_paragraph(idx))
            self.structure_table.setCellWidget(row, 2, edit_btn)
            
            # Delete button
            delete_btn = QPushButton("删除")
            delete_btn.clicked.connect(lambda checked, idx=i: self.delete_paragraph(idx))
            delete_btn.setStyleSheet("QPushButton { background-color: #dc3545; color: white; }")
            self.structure_table.setCellWidget(row, 3, delete_btn)
        
        # Add tables
        for i, table in enumerate(self.document.tables):
            row = self.structure_table.rowCount()
            self.structure_table.insertRow(row)
            
            content = f"表格 ({len(table.rows)}行 x {len(table.columns)}列)"
            
            self.structure_table.setItem(row, 0, QTableWidgetItem("表格"))
            self.structure_table.setItem(row, 1, QTableWidgetItem(content))
            
            # Edit button
            edit_btn = QPushButton("编辑")
            edit_btn.clicked.connect(lambda checked, idx=i: self.edit_table(idx))
            self.structure_table.setCellWidget(row, 2, edit_btn)
            
            # Delete button
            delete_btn = QPushButton("删除")
            delete_btn.clicked.connect(lambda checked, idx=i: self.delete_table(idx))
            delete_btn.setStyleSheet("QPushButton { background-color: #dc3545; color: white; }")
            self.structure_table.setCellWidget(row, 3, delete_btn)
    
    def edit_paragraph(self, paragraph_index: int):
        """Edit a specific paragraph"""
        if not self.document or paragraph_index >= len(self.document.paragraphs):
            return
        
        paragraph = self.document.paragraphs[paragraph_index]
        self.content_editor.setPlainText(paragraph.text)
        self.current_edit_element = ('paragraph', paragraph_index)
        self.update_content_btn.setEnabled(True)
    
    def edit_table(self, table_index: int):
        """Edit a specific table"""
        if not self.document or table_index >= len(self.document.tables):
            return
        
        table = self.document.tables[table_index]
        
        # Convert table to editable text format
        table_text = ""
        for row in table.rows:
            row_texts = [cell.text for cell in row.cells]
            table_text += "\t".join(row_texts) + "\n"
        
        self.content_editor.setPlainText(table_text)
        self.current_edit_element = ('table', table_index)
        self.update_content_btn.setEnabled(True)
    
    def update_selected_content(self):
        """Update the selected content element"""
        if not hasattr(self, 'current_edit_element'):
            return
        
        new_content = self.content_editor.toPlainText()
        element_type, element_index = self.current_edit_element
        
        try:
            # Apply variable substitution before updating
            substituted_content = self.variable_manager.substitute_variables(new_content)
            
            if element_type == 'paragraph':
                paragraph = self.document.paragraphs[element_index]
                paragraph.clear()
                paragraph.add_run(substituted_content)
                
            elif element_type == 'table':
                table = self.document.tables[element_index]
                lines = substituted_content.strip().split('\n')
                
                for row_idx, line in enumerate(lines):
                    if row_idx < len(table.rows):
                        cells = line.split('\t')
                        for col_idx, cell_text in enumerate(cells):
                            if col_idx < len(table.columns):
                                table.cell(row_idx, col_idx).text = cell_text.strip()
            
            # Refresh structure and preview
            self.populate_structure_table()
            self.refresh_preview()
            
            QMessageBox.information(self, "成功", "内容已更新")
            
        except Exception as e:
            QMessageBox.critical(self, "错误", f"更新内容失败:\n{str(e)}")
    
    def delete_paragraph(self, paragraph_index: int):
        """Delete a specific paragraph"""
        if not self.document or paragraph_index >= len(self.document.paragraphs):
            return
        
        # Confirm deletion
        reply = QMessageBox.question(
            self, "确认删除", 
            f"确定要删除第 {paragraph_index + 1} 个段落吗？",
            QMessageBox.Yes | QMessageBox.No
        )
        
        if reply == QMessageBox.Yes:
            try:
                # Remove paragraph from document
                paragraph = self.document.paragraphs[paragraph_index]
                p = paragraph._element
                p.getparent().remove(p)
                
                # Refresh structure and preview
                self.populate_structure_table()
                self.refresh_preview()
                
                QMessageBox.information(self, "成功", "段落已删除")
                
            except Exception as e:
                QMessageBox.critical(self, "错误", f"删除段落失败:\n{str(e)}")
    
    def delete_table(self, table_index: int):
        """Delete a specific table"""
        if not self.document or table_index >= len(self.document.tables):
            return
        
        # Confirm deletion
        reply = QMessageBox.question(
            self, "确认删除", 
            f"确定要删除第 {table_index + 1} 个表格吗？",
            QMessageBox.Yes | QMessageBox.No
        )
        
        if reply == QMessageBox.Yes:
            try:
                # Remove table from document
                table = self.document.tables[table_index]
                tbl = table._element
                tbl.getparent().remove(tbl)
                
                # Refresh structure and preview
                self.populate_structure_table()
                self.refresh_preview()
                
                QMessageBox.information(self, "成功", "表格已删除")
                
            except Exception as e:
                QMessageBox.critical(self, "错误", f"删除表格失败:\n{str(e)}")
    
    def refresh_variables_list(self):
        """Refresh the list of available variables"""
        self.variables_list.clear()
        
        variables = self.variable_manager.get_all_variables_with_metadata()
        
        # Always show all possible variable patterns, regardless of whether they exist
        self._add_all_possible_variables(variables)
    
    def _add_all_possible_variables(self, existing_variables):
        """Add all possible variable patterns to the list"""
        
        # Define all possible variable patterns
        variable_patterns = {
            '📅 时间变量': [
                ('patrol_time_任务名', '巡检执行时间 (HH:MM:SS)', 'time'),
                ('patrol_time_formatted_任务名', '完整格式时间 (年月日 时分秒)', 'time'), 
                ('patrol_date_任务名', '巡检执行日期 (YYYY-MM-DD)', 'date'),
                ('patrol_datetime_任务名', '日期时间 (YYYY-MM-DD HH:MM:SS)', 'datetime'),
            ],
            '📸 页面截图': [
                ('screenshot_任务名_网站', '页面整体截图', 'image'),
                ('visual_检查名_任务名', '视觉检查截图 (特定元素)', 'image'),
            ],
            '📝 提取内容': [
                ('extracted_检查名_任务名', 'XPath/CSS选择器提取的文本内容', 'text'),
                ('api_response_检查名_任务名', 'API检查返回的响应内容', 'text'),
            ],
            '✅ 状态信息': [
                ('status_检查名_任务名', '检查项状态 (成功/失败)', 'text'),
                ('status_任务名_网站', '网站巡检状态 (成功/失败)', 'text'),
                ('api_status_检查名_任务名', 'API状态码 (200, 404等)', 'number'),
                ('response_time_任务名_网站', '网站响应时间 (毫秒)', 'number'),
            ],
            '📁 下载文件': [
                ('download_path_检查名_任务名', '下载文件的完整路径', 'file'),
                ('download_name_检查名_任务名', '下载文件的文件名', 'text'),
                ('download_size_检查名_任务名', '下载文件大小 (字节)', 'number'),
            ],
            '🔗 表单检查': [
                ('form_result_检查名_任务名', '表单提交结果', 'text'),
                ('form_response_检查名_任务名', '表单提交后的响应', 'text'),
            ]
        }
        
        # Track which variables actually exist
        existing_var_names = set(existing_variables.keys()) if existing_variables else set()
        
        # Add each category
        for category_name, var_list in variable_patterns.items():
            # Add category header
            header_item = QListWidgetItem(f"━━━ {category_name} ━━━")
            header_item.setData(Qt.UserRole, None)
            header_item.setFlags(header_item.flags() & ~Qt.ItemIsSelectable)
            header_font = QFont()
            header_font.setBold(True)
            header_item.setFont(header_font)
            self.variables_list.addItem(header_item)
            
            # Add variables in this category
            for var_pattern, description, var_type in var_list:
                # Check if this variable pattern matches any existing variables
                matching_vars = [name for name in existing_var_names 
                               if self._matches_pattern(name, var_pattern)]
                
                if matching_vars:
                    # Show actual existing variables
                    for var_name in matching_vars:
                        var_info = existing_variables[var_name]
                        actual_value = var_info.get('value', '')
                        actual_description = var_info.get('metadata', {}).get('description', description)
                        
                        # Create display text with actual value preview
                        if var_type == 'image' and actual_value:
                            display_text = f"${{{var_name}}} ✅\n  📸 {actual_description}\n  📄 文件: {Path(str(actual_value)).name}"
                        elif var_type == 'text' and actual_value:
                            preview = str(actual_value)[:50] + "..." if len(str(actual_value)) > 50 else str(actual_value)
                            display_text = f"${{{var_name}}} ✅\n  📝 {actual_description}\n  💬 内容: {preview}"
                        elif var_type in ['number', 'time', 'date', 'datetime'] and actual_value:
                            display_text = f"${{{var_name}}} ✅\n  🔢 {actual_description}\n  📊 值: {actual_value}"
                        else:
                            display_text = f"${{{var_name}}} ✅\n  ℹ️ {actual_description}"
                            
                        item = QListWidgetItem(display_text)
                        item.setData(Qt.UserRole, var_name)
                        item.setToolTip(f"变量名: {var_name}\n类型: {var_type}\n描述: {actual_description}\n当前值: {actual_value}\n双击插入到编辑器")
                        
                        # Style existing variables differently
                        item.setBackground(QBrush(QColor(240, 255, 240)))  # Light green background
                        self.variables_list.addItem(item)
                else:
                    # Show pattern template for variables that don't exist yet
                    display_text = f"${{{var_pattern}}} ⏳\n  ℹ️ {description}\n  💡 执行巡检后自动生成"
                    
                    item = QListWidgetItem(display_text)
                    item.setData(Qt.UserRole, var_pattern)
                    item.setToolTip(f"变量模式: {var_pattern}\n类型: {var_type}\n描述: {description}\n状态: 将在执行巡检任务后自动生成\n双击插入模式到编辑器")
                    
                    # Style template variables differently  
                    item.setBackground(QBrush(QColor(250, 250, 250)))  # Light gray background
                    font = QFont()
                    font.setItalic(True)
                    item.setFont(font)
                    self.variables_list.addItem(item)
    
    def _matches_pattern(self, var_name, pattern):
        """Check if a variable name matches a pattern like 'patrol_time_任务名'"""
        # Convert pattern to regex-like matching
        # Since variable names can contain multiple underscores from URL cleaning,
        # we need more flexible patterns
        import re
        
        # Replace common placeholders with regex patterns
        regex_pattern = pattern
        regex_pattern = regex_pattern.replace('任务名', '__TASK_NAME__')
        regex_pattern = regex_pattern.replace('网站', '__WEBSITE__')  
        regex_pattern = regex_pattern.replace('检查名', '__CHECK_NAME__')
        
        # Escape special regex characters
        regex_pattern = re.escape(regex_pattern)
        
        # Replace our placeholders with actual regex patterns
        # Use more flexible patterns that allow underscores within values
        regex_pattern = regex_pattern.replace('__TASK_NAME__', r'[^_]+(?:_[^_]+)*')  # Allow underscores within task names
        regex_pattern = regex_pattern.replace('__WEBSITE__', r'.+')  # Websites can contain any chars after processing  
        regex_pattern = regex_pattern.replace('__CHECK_NAME__', r'[^_]+(?:_[^_]+)*')  # Allow underscores within check names
        
        # Add anchors
        regex_pattern = f"^{regex_pattern}$"
        
        try:
            return bool(re.match(regex_pattern, var_name))
        except:
            # Fallback to simple prefix matching
            pattern_prefix = pattern.split('_')[0]
            return var_name.startswith(pattern_prefix)
    
    def show_variable_selector(self):
        """Show variable selection dialog"""
        # Refresh variables list first
        self.refresh_variables_list()
        
        variables = self.variable_manager.get_all_variables_with_metadata()
        
        if not variables:
            QMessageBox.information(
                self, "提示", 
                "当前没有可用的变量。\n\n变量会在执行巡检任务时自动生成，包括:\n" +
                "• ${patrol_time_任务名} - 巡检执行时间 (HH:MM:SS)\n" +
                "• ${patrol_time_formatted_任务名} - 完整格式时间 (年月日 时分秒)\n" +
                "• ${patrol_date_任务名} - 巡检执行日期 (YYYY-MM-DD)\n" +
                "• ${patrol_datetime_任务名} - 日期时间 (YYYY-MM-DD HH:MM:SS)\n" +
                "• ${screenshot_任务名_网站} - 页面截图\n" +
                "• ${extracted_检查名_任务名} - XPath/CSS提取的内容\n" +
                "• ${status_任务名_网站} - 巡检状态 (成功/失败)\n" +
                "• ${response_time_任务名_网站} - 页面响应时间\n" +
                "• ${api_status_检查名_任务名} - API状态码\n\n" +
                "执行巡检任务后，这些变量将自动出现在右侧变量列表中。"
            )
            return
        
        QMessageBox.information(
            self, "使用变量", 
            f"发现 {len(variables)} 个可用变量！\n\n" +
            "使用方法:\n" +
            "1. 在右侧变量列表中选择需要的变量\n" +
            "2. 双击变量名即可插入到编辑器中\n" +
            "3. 变量语法: ${变量名}\n" +
            "4. 报告生成时会自动替换为实际值\n\n" +
            "提示: 将鼠标悬停在变量上可查看详细信息"
        )
    
    def insert_selected_variable(self, item):
        """Insert selected variable into the content editor"""
        var_name = item.data(Qt.UserRole)
        if var_name:  # Only insert if it's actually a variable (not a header or help text)
            cursor = self.content_editor.textCursor()
            cursor.insertText(f"${{{var_name}}}")
            self.logger.info(f"Inserted variable: ${{{var_name}}}")
    
    def load_variables_from_patrol_results(self, patrol_results):
        """Load variables from patrol results for use in reports"""
        try:
            for result in patrol_results:
                # Store screenshot as variable
                if result.screenshot_path:
                    var_name = f"screenshot_{result.task_name}_{result.website_url}".replace("://", "_").replace("/", "_").replace(".", "_")
                    self.variable_manager.set_variable(
                        var_name, 
                        result.screenshot_path, 
                        "image", 
                        f"Screenshot from {result.website_url}"
                    )
                
                # Store extracted values as variables
                if result.check_results:
                    for check_name, check_result in result.check_results.items():
                        if check_result.get('extracted_value'):
                            var_name = f"extracted_{check_name}_{result.task_name}".replace(" ", "_")
                            self.variable_manager.set_variable(
                                var_name,
                                check_result['extracted_value'],
                                "text",
                                f"Extracted value from {check_name} check"
                            )
                
                # Store other extracted data
                if result.extracted_data:
                    for key, value in result.extracted_data.items():
                        var_name = f"data_{key}_{result.task_name}".replace(" ", "_")
                        self.variable_manager.set_variable(
                            var_name,
                            value,
                            "auto",
                            f"Data extracted from {result.website_url}"
                        )
            
            # Refresh the variables list
            self.refresh_variables_list()
            
        except Exception as e:
            self.logger.error(f"Failed to load variables from patrol results: {e}")
    
    def set_variable_manager(self, variable_manager: VariableManager):
        """Set an external variable manager"""
        self.variable_manager = variable_manager
        self.refresh_variables_list()
    
    def add_paragraph(self):
        """Add a new paragraph to the document"""
        if not self.document:
            return
        
        new_text = self.content_editor.toPlainText()
        if not new_text.strip():
            QMessageBox.warning(self, "警告", "请在编辑框中输入要添加的内容")
            return
        
        try:
            # Apply variable substitution
            substituted_text = self.variable_manager.substitute_variables(new_text)
            
            # Add new paragraph at the end
            self.document.add_paragraph(substituted_text)
            
            # Refresh structure and preview
            self.populate_structure_table()
            self.refresh_preview()
            self.content_editor.clear()
            
            QMessageBox.information(self, "成功", "已添加新段落")
            
        except Exception as e:
            QMessageBox.critical(self, "错误", f"添加段落失败:\n{str(e)}")
    
    def refresh_preview(self):
        """Refresh the document preview"""
        if not self.document:
            return
        
        preview_content = "文档预览 (变量已替换):\n\n"
        
        for paragraph in self.document.paragraphs:
            if paragraph.style.name.startswith('Heading'):
                level = paragraph.style.name[-1] if paragraph.style.name[-1].isdigit() else "1"
                prefix = "#" * int(level) + " "
            else:
                prefix = ""
            
            # Apply variable substitution to preview
            paragraph_text = self.variable_manager.substitute_variables(paragraph.text)
            preview_content += prefix + paragraph_text + "\n\n"
        
        # Add tables
        for i, table in enumerate(self.document.tables):
            preview_content += f"[表格 {i+1}]\n"
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    # Apply variable substitution to table cells
                    cell_text = self.variable_manager.substitute_variables(cell.text)
                    row_texts.append(cell_text)
                row_text = " | ".join(row_texts)
                preview_content += row_text + "\n"
            preview_content += "\n"
        
        self.preview_text.setPlainText(preview_content)
    
    def save_report(self):
        """Save the current report"""
        if not self.document or not self.current_report_path:
            return
        
        try:
            self.document.save(self.current_report_path)
            QMessageBox.information(self, "成功", "报告已保存")
            self.report_saved.emit(self.current_report_path)
            
        except Exception as e:
            QMessageBox.critical(self, "错误", f"保存失败:\n{str(e)}")
    
    def save_as_report(self):
        """Save the report as a new file"""
        if not self.document:
            return
        
        file_path, _ = QFileDialog.getSaveFileName(
            self, "另存为", "", "Word文档 (*.docx);;所有文件 (*.*)"
        )
        
        if file_path:
            try:
                self.document.save(file_path)
                self.current_report_path = file_path
                QMessageBox.information(self, "成功", f"报告已保存到:\n{file_path}")
                self.report_saved.emit(file_path)
                
            except Exception as e:
                QMessageBox.critical(self, "错误", f"保存失败:\n{str(e)}")
    
    def on_structure_selection_changed(self, current_row, current_col, previous_row, previous_col):
        """Handle structure table selection changes"""
        # Could be used to auto-load content for editing
        pass